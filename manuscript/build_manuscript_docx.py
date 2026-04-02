#!/usr/bin/env python3
"""Build the manuscript .docx with embedded figures."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from pathlib import Path
import re

PROJ = Path(__file__).resolve().parent.parent
FIG_DIR = PROJ / "figures"
OUT = PROJ / "manuscript" / "Surprise_EEG_Manuscript_with_Figures.docx"

doc = Document()

# ── Page setup ──────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Styles ──────────────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(12)
style_normal.paragraph_format.space_after = Pt(6)
style_normal.paragraph_format.line_spacing = 1.5

# ── Helper functions ────────────────────────────────────────
def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(12)
    return p

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_body(text):
    p = doc.add_paragraph()
    # Handle inline bold (**...**)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        else:
            run = p.add_run(part)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    p.paragraph_format.first_line_indent = Cm(1.27)
    return p

def add_body_no_indent(text):
    p = add_body(text)
    p.paragraph_format.first_line_indent = Cm(0)
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        else:
            run = p.add_run(part)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    return p

def add_figure(filename, caption, width=Inches(6.0)):
    fpath = FIG_DIR / filename
    if not fpath.exists():
        add_body_no_indent(f"[Figure not found: {filename}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(fpath), width=width)
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_label = cap.add_run(caption.split('.')[0] + '. ')
    run_label.bold = True
    run_label.font.name = 'Times New Roman'
    run_label.font.size = Pt(10)
    rest = '.'.join(caption.split('.')[1:]).strip()
    if rest:
        run_rest = cap.add_run(rest)
        run_rest.font.name = 'Times New Roman'
        run_rest.font.size = Pt(10)
    cap.paragraph_format.space_after = Pt(12)

def add_spacer():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

def add_table(headers, rows, caption=""):
    """Add a formatted table with caption."""
    if caption:
        cap = doc.add_paragraph()
        # Split caption at first period for bold label
        if '.' in caption:
            label, rest = caption.split('.', 1)
            run_label = cap.add_run(label + '.')
            run_label.bold = True
            run_label.font.name = 'Times New Roman'
            run_label.font.size = Pt(10)
            if rest.strip():
                run_rest = cap.add_run(' ' + rest.strip())
                run_rest.font.name = 'Times New Roman'
                run_rest.font.size = Pt(10)
        else:
            run = cap.add_run(caption)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
        cap.paragraph_format.space_after = Pt(4)

    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Light Shading'

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for par in cell.paragraphs:
            for run in par.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(9)

    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for par in cell.paragraphs:
                for run in par.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(9)

    add_spacer()

# ═══════════════════════════════════════════════════════════
#  MANUSCRIPT
# ═══════════════════════════════════════════════════════════

# ── Running head ───────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("Running head: SURPRISE SIGNATURES IN HUMAN EEG")
run.font.name = 'Times New Roman'
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(128, 128, 128)
p.paragraph_format.space_after = Pt(24)

# ── Title ───────────────────────────────────────────────────
add_title("Stochastic Surprise Signatures in Human EEG:\nA Reproducible Benchmark of Prediction-Error Models\nUsing the ERP CORE Oddball Dataset")

# ── Authors and affiliations ──────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Brhanu F. Znabu")
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run2 = p.add_run("1*")
run2.font.name = 'Times New Roman'
run2.font.size = Pt(10)
run2.font.superscript = True
run3 = p.add_run(", Zohaib Atif")
run3.font.name = 'Times New Roman'
run3.font.size = Pt(12)
run4 = p.add_run("2")
run4.font.name = 'Times New Roman'
run4.font.size = Pt(10)
run4.font.superscript = True
run5 = p.add_run(", Pradeep Devkota")
run5.font.name = 'Times New Roman'
run5.font.size = Pt(12)
run6 = p.add_run("3")
run6.font.name = 'Times New Roman'
run6.font.size = Pt(10)
run6.font.superscript = True

aff1 = doc.add_paragraph()
aff1.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = aff1.add_run(
    "\u00b9Biomedical Engineering Program, College of Engineering, "
    "University of Nebraska-Lincoln, Lincoln, NE 68588-0642, USA."
)
r.font.name = 'Times New Roman'
r.font.size = Pt(10)

aff2 = doc.add_paragraph()
aff2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = aff2.add_run(
    "\u00b2Department of Biomedical Science and Engineering, Gwangju 61005, South Korea"
)
r.font.name = 'Times New Roman'
r.font.size = Pt(10)

aff3 = doc.add_paragraph()
aff3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = aff3.add_run(
    "\u00b3School of Biomedical Science, University of Kansas Medical Center, "
    "Kansas City, KS 66160, USA"
)
r.font.name = 'Times New Roman'
r.font.size = Pt(10)

add_spacer()
corr = doc.add_paragraph()
corr.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = corr.add_run(
    "Corresponding author: Brhanu F. Znabu, 25674436@nebraska.edu, "
    "ORCID: 0009-0009-7230-754X"
)
r.font.name = 'Times New Roman'
r.font.size = Pt(10)

add_spacer()
kw = doc.add_paragraph()
r = kw.add_run("Keywords: ")
r.bold = True
r.font.name = 'Times New Roman'
r.font.size = Pt(10)
r2 = kw.add_run(
    "prediction error, surprise, mismatch negativity, P3b, Bayesian inference, EEG"
)
r2.font.name = 'Times New Roman'
r2.font.size = Pt(10)

add_spacer()

# ── Abstract ────────────────────────────────────────────────
add_heading("Abstract", level=1)

add_body_no_indent(
    "The brain continuously generates predictions about incoming sensory input "
    "and produces characteristic neural responses when those predictions are violated. "
    "In EEG oddball paradigms, these prediction-error responses manifest as the "
    "mismatch negativity (MMN) and P3b components. However, \u201csurprise\u201d can be "
    "formalized in multiple ways, and no prior study has systematically compared "
    "these formulations on the same dataset using single-trial methods. Here, we "
    "implemented four hierarchical surprise estimators and applied them to the "
    "ERP CORE dataset (N = 40; N = 38 analyzed for MMN, N = 36 for P3 after "
    "exclusions; auditory MMN and visual P3 paradigms). Using linear mixed-effects "
    "encoding models, we found that Bayesian surprise, quantifying the magnitude "
    "of belief revision, showed the strongest association with single-trial MMN "
    "amplitude among all models tested (uncorrected p = 0.015), though this did "
    "not survive Holm\u2013Bonferroni correction for multiple comparisons "
    "(p_corrected = 0.062). In the P3 paradigm, change-point surprise "
    "(uncorrected p = 0.017, p_corrected = 0.069) and Bayesian surprise "
    "(uncorrected p = 0.044, p_corrected = 0.131) showed trends in the P3b "
    "window but likewise did not survive correction. High multicollinearity "
    "among Shannon-based and change-point surprise regressors "
    "(VIF > 70) limited the interpretability of direct model comparisons. All "
    "partial R\u00b2 values were extremely small (<0.02%), and cross-validated "
    "prediction showed no real gain over stimulus class alone. In a decoding "
    "analysis, contextual surprise features (residualized against stimulus type) "
    "did not improve cross-subject classification of stimulus class beyond ERP "
    "amplitude features alone (\u0394AUC = \u22120.093). We conclude that Bayesian "
    "surprise shows the strongest trend among competing models, but that "
    "stationary oddball paradigms lack the model identifiability needed to "
    "definitively distinguish surprise formulations due to high regressor "
    "multicollinearity. All data, code, and analysis pipelines "
    "are publicly available."
)

# ── 1  Introduction ─────────────────────────────────────────
add_heading("1. Introduction", level=1)

add_body(
    "The brain has been characterized as a prediction machine that continuously "
    "generates expectations about incoming sensory input and updates internal models "
    "when those expectations are violated (Rao & Ballard, 1999; Friston, 2005). This "
    "predictive processing framework provides a unifying account of perception, "
    "attention, and learning, proposing that neural activity primarily encodes the "
    "discrepancy between predicted and observed events \u2014 the prediction error."
)

add_body(
    "Oddball paradigms offer the cleanest experimental window into neural "
    "prediction-error processing. In these paradigms, rare deviant stimuli are "
    "embedded among frequent standard stimuli, and the brain produces characteristic "
    "electrophysiological responses to the deviants. The mismatch negativity (MMN) "
    "is an early fronto-central ERP component (100\u2013250 ms post-stimulus) reflecting "
    "automatic detection of auditory deviance, thought to arise from prediction-error "
    "signals in auditory cortex (N\u00e4\u00e4t\u00e4nen et al., 2007). The P3b is a later "
    "parietal component (250\u2013500 ms) reflecting context updating, attention "
    "allocation, and the conscious evaluation of surprising events (Polich, 2007)."
)

add_body(
    "While these ERP components are well-established as neural signatures of "
    "prediction error, \u201csurprise\u201d itself can be mathematically defined in multiple "
    "principled ways that make different assumptions about the underlying computation. "
    "Shannon surprise (S = \u2212log p(x)) quantifies how rare an event is based on its "
    "probability, treating surprise as a simple function of frequency. Bayesian "
    "surprise, defined as the KL divergence between posterior beliefs before and "
    "after an observation (Itti & Baldi, 2009), measures the magnitude of belief "
    "revision \u2014 how much an observer\u2019s model of the world had to change. "
    "Change-point models (Adams & MacKay, 2007) compute the posterior probability "
    "that the underlying generative process has changed, capturing surprise about "
    "regime shifts rather than individual events."
)

add_body(
    "These formulations are not merely theoretical alternatives \u2014 they make "
    "different predictions about trial-by-trial neural responses. Shannon surprise "
    "predicts that neural responses scale with stimulus rarity. Bayesian surprise "
    "predicts that responses scale with the magnitude of belief updating, which "
    "depends on the observer\u2019s uncertainty. Change-point surprise predicts maximal "
    "responses when the observer infers that the statistical environment has "
    "fundamentally changed. Prior work has examined these models individually: "
    "Mars et al. (2008) demonstrated trial-by-trial correlations between Bayesian "
    "surprise and ERP amplitudes; Kolossa et al. (2015) used model-based approaches "
    "to explain P300 fluctuations; Ostwald et al. (2012) provided evidence for "
    "neural encoding of Bayesian surprise in somatosensation. However, no study has "
    "systematically compared all of these formulations on the same dataset using "
    "single-trial methods and formal model comparison."
)

add_body(
    "Here, we address this gap by implementing four hierarchical surprise "
    "estimators \u2014 static Shannon, adaptive Shannon, Bayesian, and change-point "
    "surprise \u2014 and benchmarking their ability to explain trial-by-trial EEG "
    "prediction-error responses in the ERP CORE dataset (Kappenman et al., 2021). "
    "We use the auditory MMN paradigm as our primary target and the visual P3 "
    "paradigm as an internal replication. Our central hypotheses are: (H1) that "
    "adaptive surprise estimators, particularly Bayesian surprise, will better "
    "explain single-trial ERP responses than static frequency-based surprise; and "
    "(H2) that surprise-derived features will improve single-trial decoding of "
    "stimulus class beyond classical ERP amplitude features."
)

# ── 2  Methods ──────────────────────────────────────────────
add_heading("2. Methods", level=1)

add_heading("2.1 Dataset", level=2)
add_body(
    "We used the ERP CORE dataset (Kappenman et al., 2021), a standardized resource "
    "providing EEG data from 40 participants recorded during six optimized ERP "
    "paradigms. Sub-012 was excluded from the original release; sub-007 was excluded "
    "for excessive artifact rejection (85.4% rejection rate), yielding N = 38 for "
    "MMN analyses. We analyzed the MMN paradigm (auditory oddball with frequency "
    "deviants; ~80% standards, ~20% deviants) and the P3 paradigm (visual oddball "
    "with letter targets; ~80% non-targets, ~20% targets; sub-003, sub-005, sub-032 "
    "excluded for >80% rejection rates, yielding N = 36). Data were recorded at "
    "1024 Hz from 30 EEG channels plus 3 EOG channels using BioSemi ActiveTwo systems."
)

add_heading("2.2 Preprocessing", level=2)
add_body_no_indent(
    "All preprocessing was performed using MNE-Python 1.8.0 (Gramfort et al., 2013) "
    "with the following fixed parameters:"
)
add_bullet("**Filtering:** 0.1\u201330 Hz bandpass (zero-phase FIR, Hamming window)")
add_bullet("**Resampling:** Downsampled to 256 Hz")
add_bullet("**Reference:** Average reference")
add_bullet(
    "**Artifact rejection:** ICA (FastICA, 15 components) for eye artifact removal, "
    "with automatic EOG component detection; amplitude-based epoch rejection "
    "(\u00b1150 \u00b5V threshold) as fallback when channel positions were unavailable for autoreject"
)
add_bullet("**Epoching:** \u2212200 to 800 ms relative to stimulus onset")
add_bullet("**Baseline correction:** \u2212200 to 0 ms")

# ── Table 1: Preprocessing parameters ──
add_table(
    headers=["Parameter", "Specification"],
    rows=[
        ["High-pass filter", "0.1 Hz (zero-phase FIR, Hamming window)"],
        ["Low-pass filter", "30 Hz (zero-phase FIR, Hamming window)"],
        ["Sampling rate", "256 Hz (downsampled from 1024 Hz)"],
        ["Reference", "Average reference"],
        ["Artifact rejection (ICA)", "FastICA, 15 components, automatic EOG detection"],
        ["Artifact rejection (epochs)", "Amplitude threshold: \u00b1150 \u00b5V"],
        ["Epoch window", "\u2212200 to 800 ms (stimulus-locked)"],
        ["Baseline correction", "\u2212200 to 0 ms"],
        ["Exclusion criterion",
         "<100 epochs (MMN) or <30 epochs (P3), or >80% rejection rate"],
    ],
    caption="Table 1. Preprocessing parameters and exclusion criteria."
)

add_body(
    "After preprocessing and exclusion, the MMN dataset comprised 35,332 epochs "
    "across 38 subjects (mean 930 \u00b1 48 per subject; 7.0% rejection rate). The "
    "P3 dataset comprised 5,865 epochs across 36 subjects."
)

add_heading("2.3 Surprise Estimators", level=2)
add_body_no_indent(
    "All estimators operated on the binary stimulus sequence extracted from retained "
    "epochs (0 = standard, 1 = deviant):"
)
add_body_no_indent(
    "1. **Static Shannon surprise:** S_t = \u2212log\u2082 p_global(x_t), where p_global is "
    "the overall frequency of each stimulus type across the entire sequence. This "
    "produces a fixed value per stimulus class."
)
add_body_no_indent(
    "2. **Adaptive Shannon surprise:** S_t = \u2212log\u2082 p_w(x_t), where p_w is estimated "
    "from a sliding window of the preceding w trials (w = 20 for primary analysis; "
    "w = 10, 50 for sensitivity). Laplace smoothing was applied to prevent log(0)."
)
add_body_no_indent(
    "3. **Bayesian surprise:** D_KL[Beta(\u03b1_t, \u03b2_t) || Beta(\u03b1_{t\u22121}, \u03b2_{t\u22121})], "
    "the KL divergence between successive posteriors of a Beta-Bernoulli conjugate "
    "model with flat prior (Beta(1,1)). This measures the magnitude of belief "
    "revision after each observation."
)
add_body_no_indent(
    "4. **Change-point predictive surprise:** \u2212log\u2082 P(x_t | model), the negative "
    "log marginal predictive probability under the Adams & MacKay (2007) online "
    "change-point detection model with geometric hazard rate (h = 1/200; sensitivity "
    "analysis over h \u2208 {1/50, 1/100, 1/500}). This marginalizes over all possible "
    "run lengths, weighting each by its posterior probability."
)
add_body(
    "Secondary metrics included posterior entropy, estimated volatility (entropy of "
    "the run-length distribution), and mean run length. Regressors were z-scored "
    "before inclusion in encoding models."
)

add_heading("2.4 Feature Extraction", level=2)
add_body(
    "**ERP features:** Mean amplitude was extracted in two time windows for "
    "both paradigms: the MMN window (100\u2013250 ms, fronto-central ROI: Fz, FCz, "
    "Cz, FC3, FC4) and the P3b window (250\u2013500 ms, parietal ROI: Pz, CPz, P3, "
    "P4, CP1, CP2). Testing both windows in both paradigms allowed us to assess "
    "whether surprise effects were specific to the expected ERP component or "
    "extended across time windows."
)

add_heading("2.5 Encoding Analysis", level=2)
add_body(
    "**Primary analysis:** Linear mixed-effects models were fit with EEG amplitude "
    "as the dependent variable, stimulus class as a fixed effect (baseline model), "
    "and subject as a random intercept. Each surprise model was tested individually "
    "against the stimulus-class-only baseline (individual-model-vs-baseline design), "
    "avoiding multicollinearity from joint inclusion. Model comparison used AIC and "
    "likelihood ratio tests (LRT)."
)
add_body(
    "**Cross-validated prediction:** To assess out-of-sample predictive value, "
    "we performed leave-one-subject-out cross-validation: for each held-out "
    "subject, mixed-effects models were fit on the remaining subjects and used "
    "to predict the held-out subject\u2019s trial-by-trial amplitudes. Prediction "
    "MSE was compared between baseline and surprise-augmented models via paired "
    "t-tests across folds. Holm\u2013Bonferroni correction was applied within each "
    "ERP window (4 comparisons)."
)
add_body(
    "**Time-resolved regression:** At each time point, ERP amplitude (averaged "
    "across ROI channels) was regressed on each z-scored surprise regressor, pooled "
    "across subjects. Group-level significance was assessed via one-sample t-tests "
    "with cluster-based permutation correction (1000 permutations)."
)
add_body(
    "**Multicollinearity assessment:** Variance inflation factors (VIF) were "
    "computed for all four regressors entered simultaneously. VIF \u2265 10 was the "
    "threshold for declaring problematic collinearity."
)

add_heading("2.6 Decoding Analysis", level=2)
add_body("**Task:** Binary classification of standard vs. deviant trials.")
add_body(
    "**Features:** (i) ERP-only (MMN and P3b window amplitudes); (ii) ERP + "
    "contextual surprise (surprise regressors residualized against stimulus type "
    "to prevent label leakage)."
)
add_body(
    "**Evaluation:** Cross-subject leave-5-out CV was designated as the primary "
    "test of H2, because surprise regressors computed from the deterministic "
    "stimulus sequence encode stimulus identity and produce trivially perfect "
    "within-subject classification. Within-subject 5-fold stratified CV was "
    "reported as a secondary analysis."
)
add_body(
    "**Classifier:** L2-regularized logistic regression (C = 1.0) with balanced "
    "class weights, preceded by standard scaling."
)
add_body("**Metrics:** ROC-AUC, PR-AUC, balanced accuracy, Brier score.")

# ── 3  Results ──────────────────────────────────────────────
add_heading("3. Results", level=1)

add_heading("3.1 ERP Replication", level=2)
add_body(
    "Grand-average ERP waveforms replicated published ERP CORE results. In the "
    "MMN paradigm, deviant tones elicited a clear fronto-central negativity "
    "relative to standards in the 100\u2013250 ms window (MMN). In the P3 paradigm, "
    "target stimuli elicited a parietal positivity in the 250\u2013500 ms window (P3b). "
    "These results confirm that our preprocessing pipeline preserved the expected "
    "neural signatures (Figure 2)."
)

# ── Figure 2 ──
add_figure(
    "fig2_erp_replication_MMN.png",
    "Figure 2. ERP replication and quality control. "
    "(A) Grand-average ERP waveforms at fronto-central sites (MMN ROI) for standard "
    "(blue) and deviant (red) stimuli, with shaded \u00b1SEM. Gray bar marks the MMN "
    "window (100\u2013250 ms). (B) Grand-average at parietal sites (P3b ROI). "
    "(C) Difference waveforms (deviant \u2212 standard) at fronto-central and parietal "
    "sites. (D) Epoch counts per subject after artifact rejection."
)

add_heading("3.2 Surprise Regressor Properties", level=2)
add_body(
    "The four surprise models produced regressors with markedly different properties "
    "(Figure 1C). Static Shannon, adaptive Shannon, and change-point "
    "surprise were highly correlated (r = 0.95\u20130.99), reflecting their shared "
    "dependence on stimulus frequency. Bayesian surprise was substantially more "
    "distinct (r = 0.26\u20130.29 with all other models)."
)
add_body(
    "VIF analysis confirmed severe multicollinearity: static Shannon "
    "(VIF = 70.3), change-point (VIF = 105.9), and adaptive Shannon (VIF = 17.3) "
    "all exceeded the threshold of 10. Only Bayesian surprise (VIF = 1.1) was "
    "independent. This pattern was consistent across paradigms. We therefore "
    "interpret the individual-model-vs-baseline comparisons as the primary "
    "analysis and do not report combined models."
)

# ── Figure 1 ──
add_figure(
    "fig1_schematic_MMN.png",
    "Figure 1. Study schematic and surprise model hierarchy. "
    "(A) Oddball paradigm: rare deviants (red) embedded among frequent standards "
    "(blue). (B) Four surprise estimators arranged by increasing complexity. "
    "(C) Example surprise traces for one subject\u2019s stimulus sequence, z-scored "
    "for comparison. Gray lines mark deviant trials. (D) Analysis pipeline flowchart."
)

# ── Table 2A: Encoding results — MMN paradigm ──
add_table(
    headers=["Model", "Window", "\u0394AIC",
             "p (uncorr.)", "p (Holm)", "Partial R\u00b2", "\u03b2 [95% CI]"],
    rows=[
        ["Static Shannon", "MMN", "\u22121.4", "0.066", "0.131",
         "9.5\u00d710\u207b\u2075", "2.86 [\u22120.20, 5.93]"],
        ["Adaptive Shannon", "MMN", "+1.8", "0.674", "0.674",
         "5.0\u00d710\u207b\u2076", "\u22120.03 [\u22120.20, 0.13]"],
        ["Bayesian", "MMN", "\u22123.9", "0.015", "0.062",
         "1.6\u00d710\u207b\u2074", "\u22120.06 [\u22120.11, \u22120.01]"],
        ["Change-point", "MMN", "\u22122.3", "0.038", "0.115",
         "1.2\u00d710\u207b\u2074", "\u22120.41 [\u22120.81, \u22120.02]"],
        ["Static Shannon", "P3b", "+1.3", "0.394", "1.000",
         "2.0\u00d710\u207b\u2075", "1.51 [\u22121.98, 4.99]"],
        ["Adaptive Shannon", "P3b", "+1.6", "0.540", "1.000",
         "1.1\u00d710\u207b\u2075", "0.06 [\u22120.13, 0.24]"],
        ["Bayesian", "P3b", "+1.8", "0.679", "1.000",
         "4.4\u00d710\u207b\u2076", "\u22120.01 [\u22120.07, 0.05]"],
        ["Change-point", "P3b", "+1.7", "0.557", "1.000",
         "9.7\u00d710\u207b\u2076", "\u22120.13 [\u22120.58, 0.31]"],
    ],
    caption="Table 2A. Encoding model comparison \u2014 MMN paradigm (N = 38). "
    "Each model tested individually against a stimulus-class-only baseline. "
    "Holm\u2013Bonferroni correction applied within each ERP window (4 comparisons). "
    "Both the MMN window (100\u2013250 ms) and P3b window (250\u2013500 ms) were tested "
    "within the MMN paradigm to assess whether surprise effects extended beyond "
    "the primary MMN component."
)

# ── Table 2B: Encoding results — P3 paradigm ──
add_table(
    headers=["Model", "Window", "\u0394AIC",
             "p (uncorr.)", "p (Holm)", "\u03b2 [95% CI]"],
    rows=[
        ["Static Shannon", "P3b", "+1.7", "0.585", "1.000",
         "\u22120.35 [\u22121.59, 0.90]"],
        ["Adaptive Shannon", "P3b", "+0.2", "0.185", "0.369",
         "0.30 [\u22120.14, 0.74]"],
        ["Bayesian", "P3b", "\u22122.0", "0.044", "0.131",
         "0.16 [0.004, 0.32]"],
        ["Change-point", "P3b", "\u22123.7", "0.017", "0.069",
         "0.76 [0.13, 1.38]"],
    ],
    caption="Table 2B. Encoding model comparison \u2014 P3 paradigm (N = 36). "
    "Same analysis as Table 2A, applied to the visual P3 paradigm as an "
    "internal replication. Only the P3b window (250\u2013500 ms) is reported."
)

add_heading("3.3 Encoding Results (H1)", level=2)
add_body(
    "**MMN paradigm (primary).** Bayesian surprise showed the strongest "
    "association with single-trial MMN amplitude (\u0394AIC = \u22123.9, uncorrected "
    "p = 0.015), but did not survive Holm\u2013Bonferroni correction "
    "(p_corrected = 0.062). Change-point surprise was marginal "
    "before correction (\u0394AIC = \u22122.3, p = 0.038, p_corrected = 0.115). "
    "Static Shannon (p = 0.066) and adaptive Shannon (p = 0.674) did not "
    "reach significance. No surprise model significantly predicted P3b "
    "amplitude in the MMN paradigm (all p > 0.39; Table 2A). All partial R\u00b2 "
    "values were very small (<0.02%), indicating that surprise explains a "
    "negligible fraction of single-trial variance beyond stimulus class."
)
add_body(
    "**P3 paradigm (internal replication).** In the P3b window, change-point "
    "surprise (uncorrected p = 0.017, p_corrected = 0.069) and "
    "Bayesian surprise (uncorrected p = 0.044, p_corrected = 0.131) showed "
    "trends but did not survive Holm\u2013Bonferroni correction (Table 2B). The "
    "fronto-central (MMN) window was also tested in the P3 paradigm but yielded "
    "no significant results for any model (all p > 0.40; not tabulated)."
)
add_body(
    "**Cross-validated prediction.** Leave-one-subject-out cross-validation "
    "showed no significant improvement in out-of-sample prediction for any "
    "surprise model over the baseline (all p > 0.64), confirming the minimal "
    "explanatory benefit."
)
add_body(
    "**Time-resolved analysis.** All four surprise models showed significant "
    "clusters of regression coefficients in the MMN window (approximately "
    "100\u2013230 ms) at fronto-central sites. At parietal sites, significant clusters "
    "emerged in both the MMN window (125\u2013230 ms) and the P3b window "
    "(330\u2013435 ms) for all models except Bayesian surprise, which showed a later "
    "cluster (344\u2013383 ms, Figure 3)."
)
add_body(
    "**Summary for H1:** H1 is partially supported. Bayesian surprise shows "
    "the strongest trend toward predicting single-trial prediction-error "
    "responses across both paradigms, but no model survives correction for "
    "multiple comparisons. Effect sizes are very small, and the high "
    "collinearity among Shannon-based models limits the specificity of "
    "model comparisons."
)

# ── Figure 3 ──
add_figure(
    "fig3_encoding_MMN.png",
    "Figure 3. Surprise encoding results (MMN paradigm). "
    "(A) Time-resolved regression coefficients at fronto-central sites for each "
    "surprise model (colored lines, \u00b195% CI shading). Gray bar marks the MMN "
    "window. (B) \u0394AIC relative to stimulus-class baseline in the MMN window; "
    "negative = better fit. Asterisks mark uncorrected likelihood ratio tests "
    "(p < 0.05); neither survives Holm\u2013Bonferroni correction. "
    "(C) \u0394AIC in the P3b window (no model significant). "
    "(D) Pseudo-R\u00b2 improvement over baseline."
)

add_heading("3.4 Decoding Results (H2)", level=2)
add_body(
    "**Cross-subject (primary).** ERP-only features yielded a cross-subject AUC "
    "of 0.543 (MMN) and 0.604 (P3). Adding residualized contextual surprise "
    "features did not improve classification; AUC decreased to 0.450 "
    "(MMN, \u0394AUC = \u22120.093) and 0.571 (P3, \u0394AUC = \u22120.033). The decrease was "
    "marginally significant for MMN (paired t-test, p = 0.050) but not for "
    "P3 (p = 0.190)."
)
add_body(
    "**Within-subject (secondary).** Within-subject classification with surprise "
    "features yielded perfect AUC (1.000) for both paradigms, confirming the "
    "predicted label-leakage confound: surprise regressors are deterministic "
    "functions of the stimulus sequence and therefore perfectly encode stimulus "
    "identity within each subject. This result is reported for transparency but "
    "is not scientifically informative."
)
# ── Table 3: Decoding results ──
add_table(
    headers=["Feature Set", "Evaluation", "Paradigm", "ROC-AUC", "PR-AUC",
             "Balanced Acc."],
    rows=[
        ["ERP-only", "Cross-subject", "MMN", "0.543", "0.226", "0.531"],
        ["ERP+surprise (resid.)", "Cross-subject", "MMN", "0.450", "0.186",
         "0.474"],
        ["ERP-only", "Cross-subject", "P3", "0.604", "\u2014", "\u2014"],
        ["ERP+surprise (resid.)", "Cross-subject", "P3", "0.571", "\u2014",
         "\u2014"],
        ["ERP-only", "Within-subject", "MMN", "0.534 \u00b1 0.041", "\u2014",
         "\u2014"],
        ["ERP+surprise (resid.)", "Within-subject", "MMN",
         "1.000 \u00b1 0.000*", "\u2014", "\u2014"],
        ["ERP-only", "Within-subject", "P3", "0.602 \u00b1 0.128", "\u2014",
         "\u2014"],
        ["ERP+surprise (resid.)", "Within-subject", "P3",
         "1.000 \u00b1 0.000*", "\u2014", "\u2014"],
    ],
    caption="Table 3. Decoding results. "
    "*Within-subject AUC = 1.0 reflects label leakage from deterministic "
    "surprise regressors (see Section 4.3)."
)

add_body(
    "**Summary for H2:** H2 is not supported. Contextual surprise features "
    "(orthogonalized to stimulus type) do not improve cross-subject decoding "
    "and may slightly hurt performance through increased dimensionality. The "
    "within-subject analysis is uninformative due to fundamental label leakage."
)

# ── Figure 4 ──
add_figure(
    "fig4_decoding_MMN.png",
    "Figure 4. Decoding benchmark. "
    "(A) ROC curves for cross-subject classification (primary evaluation) by "
    "feature set. Dashed diagonal = chance. (B) Feature ablation: ROC-AUC by "
    "feature set, cross-subject. (C) Cross-subject vs. within-subject AUC "
    "comparison for ERP-only and ERP+surprise feature sets."
)

add_heading("3.5 Sensitivity Analyses", level=2)
add_body(
    "**Adaptive Shannon window size.** Results were qualitatively similar across "
    "window sizes (w = 10, 20, 50). The w = 20 window yielded the highest "
    "correlation with the neural data but none reached significance."
)
add_body(
    "**Change-point hazard rate.** Sensitivity analysis over "
    "h \u2208 {1/50, 1/100, 1/200, 1/500} showed that the change-point surprise "
    "regressor was highly correlated with static Shannon across all "
    "hazard rates (r > 0.97), consistent with the stationary nature of the "
    "oddball paradigm. The hazard rate had minimal impact on the encoding results."
)

# ── 4  Discussion ───────────────────────────────────────────
add_heading("4. Discussion", level=1)

add_heading("4.1 Key Findings", level=2)
add_body(
    "This study provides the first systematic, single-trial comparison of four "
    "hierarchical surprise formulations applied to prediction-error responses in "
    "human EEG. Bayesian surprise \u2014 quantifying the magnitude of belief "
    "revision \u2014 showed the strongest trend toward predicting trial-by-trial MMN "
    "amplitude, but this did not survive correction for multiple comparisons. "
    "A similar pattern emerged in the P3 paradigm, where both Bayesian and "
    "change-point surprise showed trends for P3b amplitude. No model "
    "significantly improved out-of-sample prediction or cross-subject decoding. "
    "These results suggest that while computational models of surprise capture "
    "meaningful aspects of neural prediction error, the effect sizes in "
    "stationary oddball paradigms are too small to reliably distinguish between "
    "formulations."
)

add_heading("4.2 Relation to Prior Work", level=2)
add_body(
    "Our results are consistent with Mars et al. (2008), who reported trial-by-trial "
    "correlations between model-based surprise and ERP amplitudes, and with "
    "Ostwald et al. (2012), who found evidence for neural encoding of Bayesian "
    "surprise in somatosensation. We extend these findings by providing a formal "
    "comparison across four model families within a single framework."
)
add_body(
    "The superiority of Bayesian surprise aligns with predictive coding theories "
    "proposing that the brain maintains probabilistic generative models that are "
    "updated upon receiving new evidence (Friston, 2005). Under this view, the MMN "
    "reflects the precision-weighted prediction error that drives belief updating "
    "\u2014 a quantity most directly captured by Bayesian surprise (KL divergence "
    "between successive posteriors)."
)
add_body(
    "The relatively poor performance of adaptive Shannon surprise is notable: "
    "despite capturing local frequency adaptation, the sliding-window estimator "
    "did not significantly predict neural responses beyond what static global "
    "probability already explains. This suggests that the brain\u2019s internal model "
    "may be more sophisticated than a simple frequency counter."
)

add_heading("4.3 Methodological Considerations", level=2)
add_body(
    "**Multicollinearity.** The high correlations among Shannon-based and "
    "change-point regressors (r > 0.95, VIF > 70) represent a fundamental "
    "limitation of applying these models to stationary oddball sequences. Because "
    "the deviant probability is constant throughout the experiment, all "
    "frequency-based models converge on similar estimates, and the change-point "
    "model\u2019s predictive surprise is dominated by its frequency estimate. Bayesian "
    "surprise, which measures belief change rather than event probability, provides "
    "a qualitatively different signal and is the only regressor with acceptable "
    "VIF (1.1). Future studies using roving oddball or volatile environments would "
    "better differentiate these models."
)
add_body(
    "**Label leakage in decoding.** Our within-subject decoding analysis yielded "
    "perfect classification (AUC = 1.0) when surprise features were included, even "
    "after residualizing against stimulus type. This occurs because surprise "
    "regressors are deterministic functions of the stimulus sequence: knowing the "
    "surprise values at each trial position uniquely identifies the stimulus type. "
    "This is not a bug but a fundamental property of these models when applied to "
    "a known, fixed sequence. We designated cross-subject decoding as the primary "
    "test of H2 specifically because of this confound, and we report the "
    "within-subject results transparently to alert future researchers."
)
add_body(
    "**Effect sizes.** The effect sizes for Bayesian surprise predicting MMN "
    "amplitude are very small (partial R\u00b2 < 0.02%, \u0394AIC = \u22123.9 on a dataset "
    "of 35,332 observations). Cross-validated prediction showed no improvement "
    "over the baseline model, suggesting that the encoding-model improvements "
    "may reflect overfitting to in-sample noise rather than genuine predictive "
    "power. This is consistent with the well-known difficulty of single-trial "
    "EEG analysis."
)

add_heading("4.4 Limitations", level=2)
add_body(
    "First, the oddball paradigm is stationary by design, limiting the ability to "
    "differentiate between surprise models that depend on non-stationarity "
    "(particularly the change-point model). Second, our EEG preprocessing used "
    "amplitude-based rejection rather than the planned autoreject algorithm, due "
    "to missing electrode position information in the ERP CORE files; this may "
    "have retained some artifacts. Third, with N = 38 (MMN) and N = 36 (P3) subjects, our power to "
    "detect small differences between surprise models in pairwise comparisons is "
    "limited. Fourth, we examined only two ERP paradigms; generalization to more "
    "complex or naturalistic contexts remains to be established."
)

add_heading("4.5 Future Directions", level=2)
add_body(
    "Three extensions would strengthen these findings. First, applying the same "
    "benchmark to roving oddball paradigms, where stimulus probabilities genuinely "
    "change over time, would better differentiate the change-point model from "
    "frequency-based alternatives. Second, extending the approach to naturalistic "
    "stimuli (e.g., language processing, social prediction) would test whether "
    "Bayesian surprise remains the best predictor in more ecologically valid "
    "settings. Third, combining our computational approach with source localization "
    "or intracranial recordings could reveal whether different surprise computations "
    "are realized in distinct neural circuits."
)

# ── 5  Conclusion ───────────────────────────────────────────
add_heading("5. Conclusion", level=1)
add_body(
    "This paper reports a benchmark comparing four surprise formulations as "
    "predictors of single-trial EEG responses. Bayesian surprise shows the "
    "strongest trend, but no model survives correction for multiple comparisons. "
    "High regressor collinearity and trivial label leakage are identified as "
    "critical methodological challenges. These results set a baseline and point "
    "to non-stationary paradigms as the necessary next step."
)

# ── Data and Code Availability ──────────────────────────────
add_heading("Data and Code Availability", level=1)
add_body_no_indent(
    "Data come from the ERP CORE dataset (Kappenman et al., 2021; "
    "https://erpinfo.org/erp-core; CC BY-SA 4.0). Code is available at "
    "https://github.com/brhanufen/surprise-eeg-benchmark. The conda environment "
    "specification ensures full reproducibility."
)

# ── Ethics Statement ──────────────────────────────────────
add_heading("Ethics Statement", level=1)
add_body_no_indent(
    "This study used publicly available, de-identified data from the ERP CORE "
    "dataset. No ethical approval was required for secondary analysis of "
    "de-identified public data."
)

# ── Conflict of Interest ─────────────────────────────────
add_heading("Conflict of Interest Statement", level=1)
add_body_no_indent("The authors declare no conflicts of interest.")

# ── Funding ──────────────────────────────────────────────
add_heading("Funding", level=1)
add_body_no_indent(
    "This research received no specific grant from any funding agency in the "
    "public, commercial, or not-for-profit sectors."
)

# ── Author Contributions ─────────────────────────────────
add_heading("Author Contributions (CRediT)", level=1)
add_body_no_indent(
    "B.F.Z. conceived and designed the study, performed all data processing and "
    "statistical analyses, developed the computational pipeline, and wrote the "
    "original draft of the manuscript. Z.A. contributed to the study design and "
    "methodology. P.D. contributed to manuscript writing and editing. All authors "
    "read and approved the final manuscript."
)

# ── References ──────────────────────────────────────────────
add_heading("References", level=1)

refs = [
    "Adams, R. P., & MacKay, D. J. C. (2007). Bayesian online changepoint detection. arXiv:0710.3742.",
    "Friston, K. (2005). A theory of cortical responses. Philosophical Transactions of the Royal Society B, 360, 815\u2013836.",
    "Gramfort, A., et al. (2013). MEG and EEG data analysis with MNE-Python. Frontiers in Neuroscience, 7, 267.",
    "Itti, L., & Baldi, P. (2009). Bayesian surprise attracts human attention. Vision Research, 49(10), 1295\u20131306.",
    "Jas, M., et al. (2017). Autoreject: Automated artifact rejection for MEG and EEG data. NeuroImage, 159, 417\u2013429.",
    "Kappenman, E. S., Farrens, J. L., Zhang, W., Stewart, A. X., & Luck, S. J. (2021). ERP CORE: An open resource for human event-related potential research. NeuroImage, 225, 117465.",
    "Kolossa, A., Fingscheidt, T., Wessel, K., & Kopp, B. (2015). A model-based approach to trial-by-trial P300 amplitude fluctuations. Frontiers in Human Neuroscience, 6, 359.",
    "Maris, E., & Oostenveld, R. (2007). Nonparametric statistical testing of EEG- and MEG-data. Journal of Neuroscience Methods, 164(1), 177\u2013190.",
    "Mars, R. B., et al. (2008). Trial-by-trial fluctuations in the event-related electroencephalogram reflect dynamic changes in the degree of surprise. Journal of Neuroscience, 28(47), 12539\u201312545.",
    "N\u00e4\u00e4t\u00e4nen, R., Paavilainen, P., Rinne, T., & Alho, K. (2007). The mismatch negativity (MMN) in basic research of central auditory processing. Clinical Neurophysiology, 118(12), 2544\u20132590.",
    "Ostwald, D., et al. (2012). Evidence for neural encoding of Bayesian surprise in human somatosensation. NeuroImage, 62(1), 177\u2013188.",
    "Pernet, C. R., et al. (2011). Robust, bias-free single trial ERP analysis. NeuroImage, 55(2), 604\u2013613.",
    "Polich, J. (2007). Updating P300: An integrative theory of P3a and P3b. Clinical Neurophysiology, 118(10), 2128\u20132148.",
    "Rao, R. P. N., & Ballard, D. H. (1999). Predictive coding in the visual cortex. Nature Neuroscience, 2(1), 79\u201387.",
]
for ref in refs:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.space_after = Pt(4)

# ── Supplementary Material ────────────────────────────────
add_heading("Supplementary Material: Simulation Analysis", level=1)
add_heading("Model Distinguishability in Non-Stationary Paradigms", level=2)

add_body(
    "The stationary oddball paradigm used in our main analyses produces highly "
    "correlated surprise regressors (VIF > 70 for Shannon and change-point models), "
    "preventing meaningful model comparison. To assess whether these models can be "
    "distinguished with appropriate experimental designs, we conducted a simulation "
    "study using a roving oddball paradigm with known change points."
)

add_heading("Methods", level=3)
add_body(
    "**Stimulus sequences.** We generated non-stationary oddball sequences (800 "
    "trials, 10 blocks of 80 trials) where the deviant probability alternated "
    "between low (p = 0.10\u20130.15) and high (p = 0.35\u20130.50) blocks, creating "
    "9 change points. For comparison, we also generated stationary sequences "
    "with fixed p = 0.20."
)
add_body(
    "**Regressor dissociability.** All four surprise models were computed on each "
    "sequence type, and we compared correlation matrices and VIF values."
)
add_body(
    "**Model recovery.** For each of the four surprise models as ground truth, we "
    "simulated single-trial EEG data: amplitude = \u03b2\u2080 + \u03b2\u2081\u00b7stimulus + "
    "\u03b2\u2082\u00b7surprise_true + noise (effect size \u03b2\u2082 = 0.05\u20130.50, noise \u03c3 = 1.0). "
    "All four models were then fit to the simulated data, and the model with the "
    "lowest AIC was selected as the recovered model. We ran 500 simulations per "
    "ground-truth model at \u03b2 = 0.20 and swept across 6 effect sizes with 200 "
    "simulations each."
)

add_heading("Results", level=3)
add_body(
    "**Regressor dissociability.** In the roving paradigm, mean VIF decreased from "
    "11.7 (stationary) to 6.5 (roving), a 44% reduction. Static Shannon VIF "
    "dropped from 10.9 to 2.4 (78% reduction). However, adaptive Shannon and "
    "change-point remained moderately collinear (VIF \u2248 10\u201312) even in the roving "
    "design."
)
add_body(
    "**Model recovery.** At effect size \u03b2 = 0.20, overall recovery rate was "
    "65.2%. Critically, recovery rates differed dramatically across models:"
)
add_bullet(
    "**Bayesian surprise: 100% recovery.** The KL divergence signal is distinct "
    "from frequency-based models, producing spikes at transitions that no other "
    "model replicates."
)
add_bullet(
    "**Adaptive Shannon: 82% recovery.** The sliding window produces a delayed "
    "tracking signal that is distinguishable from Bayesian and static models."
)
add_bullet(
    "**Change-point: 79% recovery.** The predictive surprise signal is similar "
    "to adaptive Shannon in roving paradigms (r = 0.95), creating some confusion "
    "between these models."
)
add_bullet(
    "**Static Shannon: 0% recovery.** Because static Shannon uses global "
    "frequencies and does not adapt, its signal is indistinguishable from the "
    "other frequency-based models. It was consistently confused with Bayesian "
    "(45%) and change-point (28%) surprise."
)
add_body(
    "**Effect size sweep.** Bayesian surprise reached \u226594% recovery at "
    "\u03b2 = 0.10 and 100% at \u03b2 \u2265 0.15. Adaptive Shannon and change-point "
    "required \u03b2 \u2265 0.20 for >80% recovery. Static Shannon never exceeded "
    "chance regardless of effect size, indicating a fundamental identifiability "
    "problem."
)

# ── Figure 5 ──
add_figure(
    "fig5_simulation.png",
    "Figure 5. Simulation results for non-stationary (roving) oddball paradigms. "
    "(A) Surprise traces in a stationary paradigm, all models highly correlated. "
    "(B) Surprise traces in a roving paradigm, models diverge at change points "
    "(dashed lines). Dotted line shows true deviant probability. "
    "(C) VIF comparison: roving paradigms substantially reduce multicollinearity. "
    "(D) Model recovery confusion matrix (\u03b2 = 0.20, 500 simulations). Bayesian "
    "surprise is perfectly recoverable; static Shannon is never recovered. "
    "(E) Recovery rate as a function of effect size. "
    "(F) Design recommendations for future studies."
)

add_heading("Implications", level=3)
add_body(
    "The inability to distinguish surprise models in our main analyses is a "
    "property of the stationary oddball paradigm, not an inherent limitation of "
    "the computational framework. Three concrete recommendations emerge:"
)
add_body(
    "First, roving or volatile paradigms are essential for testing competing "
    "surprise theories. Fixed-probability oddballs render most model comparisons "
    "uninformative due to multicollinearity."
)
add_body(
    "Second, Bayesian surprise is the most identifiable model and should be "
    "prioritized in future benchmark studies. Its unique sensitivity to belief "
    "revision produces signals that no frequency-based model can mimic."
)
add_body(
    "Third, static Shannon surprise should not be included as a competing model "
    "in formal comparisons, as it is fundamentally unidentifiable \u2014 it cannot "
    "be recovered even under ideal conditions. Its role should be limited to "
    "serving as a naive baseline."
)

# ── Save ────────────────────────────────────────────────────
doc.save(str(OUT))
print(f"Saved to {OUT}")
